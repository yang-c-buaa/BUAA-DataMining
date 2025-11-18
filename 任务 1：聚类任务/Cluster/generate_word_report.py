#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成Word格式的聚类分析报告
"""

import os
import json
import numpy as np
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler
from sklearn.metrics import silhouette_score, adjusted_rand_score
import warnings
warnings.filterwarnings('ignore')


class WordReportGenerator:
    def __init__(self, dataset_path='dataset', labels_path='cluster_labels.json', 
                 image_path='clustering_results.png'):
        """
        初始化报告生成器
        
        Args:
            dataset_path: 图像数据集路径
            labels_path: 真实标签文件路径
            image_path: 可视化结果图片路径
        """
        self.dataset_path = dataset_path
        self.labels_path = labels_path
        self.image_path = image_path
        self.doc = Document()
        
        # 类别映射
        self.label_to_id = {
            'cable': 0, 'tile': 1, 'bottle': 2, 
            'pill': 3, 'leather': 4, 'transistor': 5
        }
        self.id_to_label = {v: k for k, v in self.label_to_id.items()}
        
    def setup_document_style(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
        
    def add_title(self, text, level=1):
        """添加标题"""
        if level == 1:
            heading = self.doc.add_heading(text, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            self.doc.add_heading(text, level=level)
        return heading if level == 1 else None
    
    def add_paragraph(self, text, bold=False, italic=False):
        """添加段落"""
        p = self.doc.add_paragraph(text)
        if bold:
            p.runs[0].bold = True
        if italic:
            p.runs[0].italic = True
        return p
    
    def add_image(self, image_path, width=Inches(6)):
        """添加图片"""
        if os.path.exists(image_path):
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(image_path, width=width)
        else:
            self.add_paragraph(f"图片文件不存在: {image_path}", italic=True)
    
    def load_and_analyze_data(self):
        """加载数据并进行分析"""
        print("正在加载数据并进行分析...")
        
        # 加载真实标签
        with open(self.labels_path, 'r', encoding='utf-8') as f:
            labels_dict = json.load(f)
        
        # 加载图像
        images = []
        true_labels = []
        image_names = []
        
        for filename in sorted(os.listdir(self.dataset_path)):
            if filename.endswith('.png'):
                image_path = os.path.join(self.dataset_path, filename)
                image = Image.open(image_path)
                images.append(image)
                image_names.append(filename)
                true_label = labels_dict[filename]
                true_labels.append(self.label_to_id[true_label])
        
        # 提取特征
        features_list = []
        for image in images:
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # RGB直方图特征
            r_hist = np.histogram(image.getchannel('R'), bins=32, range=(0, 256))[0]
            g_hist = np.histogram(image.getchannel('G'), bins=32, range=(0, 256))[0]
            b_hist = np.histogram(image.getchannel('B'), bins=32, range=(0, 256))[0]
            
            r_hist = r_hist / np.sum(r_hist)
            g_hist = g_hist / np.sum(g_hist)
            b_hist = b_hist / np.sum(b_hist)
            
            feature = np.concatenate([r_hist, g_hist, b_hist])
            features_list.append(feature)
        
        features = np.array(features_list)
        
        # PCA降维
        scaler = StandardScaler()
        features_scaled = scaler.fit_transform(features)
        pca = PCA(n_components=50)
        features_pca = pca.fit_transform(features_scaled)
        
        # K-means聚类
        kmeans = KMeans(n_clusters=6, random_state=42, n_init=10)
        cluster_labels = kmeans.fit_predict(features_pca)
        
        # 计算评估指标
        silhouette_avg = silhouette_score(features_pca, cluster_labels)
        ari = adjusted_rand_score(true_labels, cluster_labels)
        purity = self._calculate_purity(true_labels, cluster_labels)
        
        # 统计信息
        true_counts = Counter(true_labels)
        cluster_counts = Counter(cluster_labels)
        
        return {
            'total_images': len(images),
            'true_counts': true_counts,
            'cluster_counts': cluster_counts,
            'silhouette_score': silhouette_avg,
            'ari': ari,
            'purity': purity,
            'pca_variance': pca.explained_variance_ratio_.sum(),
            'feature_dim': features.shape[1],
            'pca_dim': features_pca.shape[1]
        }
    
    def _calculate_purity(self, true_labels, cluster_labels):
        """计算聚类纯度"""
        total = len(true_labels)
        correct = 0
        
        for cluster_id in np.unique(cluster_labels):
            cluster_mask = np.array(cluster_labels) == cluster_id
            cluster_true_labels = np.array(true_labels)[cluster_mask]
            
            if len(cluster_true_labels) > 0:
                most_common = Counter(cluster_true_labels).most_common(1)[0][0]
                correct += np.sum(cluster_true_labels == most_common)
        
        return correct / total
    
    def generate_report(self, output_path='聚类分析报告.docx'):
        """生成完整的Word报告"""
        print("正在生成Word报告...")
        
        # 设置文档样式
        self.setup_document_style()
        
        # 标题
        self.add_title('图像聚类分析报告', level=1)
        
        # 执行分析
        analysis_results = self.load_and_analyze_data()
        
        # 1. 项目概述
        self.add_title('1. 项目概述', level=2)
        self.add_paragraph(
            f"本项目对包含 {analysis_results['total_images']} 张图像的数据集进行聚类分析。"
            f"数据集包含6个类别：cable（电缆）、tile（瓷砖）、bottle（瓶子）、"
            f"pill（药丸）、leather（皮革）和transistor（晶体管）。"
        )
        
        # 2. 数据集信息
        self.add_title('2. 数据集信息', level=2)
        self.add_paragraph("数据集类别分布如下：", bold=True)
        
        # 创建类别分布表格
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '类别'
        hdr_cells[1].text = '类别名称'
        hdr_cells[2].text = '样本数量'
        
        for label_id in range(6):
            row_cells = table.add_row().cells
            row_cells[0].text = str(label_id)
            row_cells[1].text = self.id_to_label[label_id]
            row_cells[2].text = str(analysis_results['true_counts'].get(label_id, 0))
        
        # 3. 方法说明
        self.add_title('3. 方法说明', level=2)
        
        self.add_title('3.1 特征提取', level=3)
        self.add_paragraph(
            f"采用RGB颜色直方图方法提取图像特征。对每张图像的R、G、B三个通道分别计算32-bin直方图，"
            f"并进行归一化处理。最终特征维度为 {analysis_results['feature_dim']} 维。"
        )
        
        self.add_title('3.2 特征降维', level=3)
        self.add_paragraph(
            f"使用主成分分析（PCA）对特征进行降维，保留 {analysis_results['pca_dim']} 个主成分。"
            f"PCA累计解释方差比为 {analysis_results['pca_variance']:.3f}，"
            f"能够保留原始特征的大部分信息。"
        )
        
        self.add_title('3.3 聚类算法', level=3)
        self.add_paragraph(
            "采用K-means聚类算法，设置聚类数量为6（与真实类别数一致）。"
            "K-means是一种基于距离的划分聚类方法，通过迭代优化聚类中心来最小化类内距离。"
        )
        
        # 4. 结果分析
        self.add_title('4. 结果分析', level=2)
        
        self.add_title('4.1 评估指标', level=3)
        self.add_paragraph("聚类结果评估指标如下：", bold=True)
        
        # 创建评估指标表格
        metrics_table = self.doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Light Grid Accent 1'
        
        metrics_data = [
            ('轮廓系数 (Silhouette Score)', f"{analysis_results['silhouette_score']:.4f}"),
            ('调整兰德指数 (ARI)', f"{analysis_results['ari']:.4f}"),
            ('聚类纯度 (Purity)', f"{analysis_results['purity']:.4f}"),
        ]
        
        for i, (metric, value) in enumerate(metrics_data):
            metrics_table.rows[i].cells[0].text = metric
            metrics_table.rows[i].cells[1].text = value
        
        # 添加指标说明
        self.add_paragraph(
            "• 轮廓系数：衡量样本与其所属簇的相似度，取值范围[-1, 1]，值越大越好。"
        )
        self.add_paragraph(
            "• 调整兰德指数：衡量聚类结果与真实标签的一致性，取值范围[-1, 1]，值越大越好。"
        )
        self.add_paragraph(
            "• 聚类纯度：每个簇中最多数类别的样本占比，取值范围[0, 1]，值越大越好。"
        )
        
        self.add_title('4.2 聚类分布', level=3)
        self.add_paragraph("各聚类簇的样本数量分布如下：", bold=True)
        
        # 创建聚类分布表格
        cluster_table = self.doc.add_table(rows=1, cols=2)
        cluster_table.style = 'Light Grid Accent 1'
        hdr_cells = cluster_table.rows[0].cells
        hdr_cells[0].text = '聚类簇'
        hdr_cells[1].text = '样本数量'
        
        for cluster_id in sorted(analysis_results['cluster_counts'].keys()):
            row_cells = cluster_table.add_row().cells
            row_cells[0].text = f"簇 {cluster_id}"
            row_cells[1].text = str(analysis_results['cluster_counts'][cluster_id])
        
        # 5. 可视化结果
        self.add_title('5. 可视化结果', level=2)
        self.add_paragraph(
            "下图展示了聚类结果的可视化分析，包括真实标签分布、聚类结果分布、"
            "类别分布对比和混淆矩阵。"
        )
        
        # 插入图片
        if os.path.exists(self.image_path):
            self.add_image(self.image_path, width=Inches(6))
        else:
            self.add_paragraph("可视化结果图片未找到。", italic=True)
        
        # 6. 结论
        self.add_title('6. 结论', level=2)
        
        # 根据评估指标给出结论
        silhouette = analysis_results['silhouette_score']
        ari = analysis_results['ari']
        purity = analysis_results['purity']
        
        conclusion_text = (
            f"通过RGB颜色直方图特征提取和K-means聚类算法，成功对 {analysis_results['total_images']} 张图像进行了聚类分析。"
            f"评估结果显示：轮廓系数为 {silhouette:.4f}，"
        )
        
        if silhouette > 0.5:
            conclusion_text += "表明聚类结果具有良好的类内凝聚性和类间分离性；"
        elif silhouette > 0.3:
            conclusion_text += "表明聚类结果具有一定的合理性；"
        else:
            conclusion_text += "表明聚类结果需要进一步优化；"
        
        conclusion_text += (
            f"调整兰德指数为 {ari:.4f}，"
        )
        
        if ari > 0.5:
            conclusion_text += "说明聚类结果与真实标签具有较高的一致性；"
        elif ari > 0.3:
            conclusion_text += "说明聚类结果与真实标签具有一定的一致性；"
        else:
            conclusion_text += "说明聚类结果与真实标签的一致性有待提高；"
        
        conclusion_text += (
            f"聚类纯度为 {purity:.4f}，"
        )
        
        if purity > 0.7:
            conclusion_text += "表明大多数簇能够较好地对应单一类别。"
        elif purity > 0.5:
            conclusion_text += "表明部分簇能够较好地对应单一类别。"
        else:
            conclusion_text += "表明簇与类别的对应关系需要进一步优化。"
        
        self.add_paragraph(conclusion_text)
        
        # 保存文档
        self.doc.save(output_path)
        print(f"报告已保存到: {output_path}")
        return output_path


def main():
    """主函数"""
    generator = WordReportGenerator()
    output_file = generator.generate_report('聚类分析报告.docx')
    print(f"\nWord报告生成完成！文件位置: {output_file}")


if __name__ == "__main__":
    main()


